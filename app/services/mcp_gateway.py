from __future__ import annotations

import asyncio
import csv
import html
import hashlib
import json
import mimetypes
import os
import re
import shutil
import subprocess
import uuid
from pathlib import Path
from typing import Any

import httpx

from app import db
from app.schemas import McpServerCreate
from app.services.network_policy import env_flag, require_outbound_network, validate_outbound_http_url

BASE_DIR = Path(__file__).resolve().parents[2]
ARTIFACT_DIR = Path(os.getenv("APP_ARTIFACT_DIR", str(BASE_DIR / "data" / "artifacts")))
PRESENTATION_GENERATOR_SCRIPT = BASE_DIR / "app" / "services" / "presentation_generator.mjs"


class ToolError(RuntimeError):
    pass


MCP_SECRET_PLACEHOLDER = "__MCP_SECRET_CONFIGURED__"
_MCP_SECRET_KEY_PARTS = (
    "api_key",
    "apikey",
    "authorization",
    "auth_token",
    "cookie",
    "credential",
    "password",
    "private_key",
    "secret",
    "token",
)
_MCP_SECRET_CONTAINER_KEYS = frozenset({"env", "environment", "headers", "query", "query_params"})


def _is_mcp_secret_key(key: Any) -> bool:
    normalized = re.sub(r"[^a-z0-9]+", "_", str(key).strip().lower()).strip("_")
    return bool(normalized) and any(part in normalized for part in _MCP_SECRET_KEY_PARTS)


def _has_configured_secret(value: Any) -> bool:
    if value is None or value is False:
        return False
    if isinstance(value, str):
        return bool(value.strip()) and value != MCP_SECRET_PLACEHOLDER
    if isinstance(value, (dict, list, tuple)):
        return bool(value)
    return True


def _redact_mcp_config(value: Any, *, secret_container: bool = False) -> tuple[Any, bool]:
    """Return an editable public config plus a recursive secret-presence flag."""

    if isinstance(value, dict):
        safe: dict[str, Any] = {}
        has_secret = False
        for raw_key, item in value.items():
            key = str(raw_key)
            normalized_key = re.sub(r"[^a-z0-9]+", "_", key.strip().lower()).strip("_")
            if secret_container or _is_mcp_secret_key(key):
                configured = _has_configured_secret(item)
                safe[key] = MCP_SECRET_PLACEHOLDER if configured else item
                has_secret = has_secret or configured
                continue
            redacted, nested_has_secret = _redact_mcp_config(
                item, secret_container=normalized_key in _MCP_SECRET_CONTAINER_KEYS
            )
            safe[key] = redacted
            has_secret = has_secret or nested_has_secret
        return safe, has_secret
    if isinstance(value, list):
        safe_items = []
        has_secret = False
        for item in value:
            redacted, nested_has_secret = _redact_mcp_config(
                item, secret_container=secret_container
            )
            safe_items.append(redacted)
            has_secret = has_secret or nested_has_secret
        return safe_items, has_secret
    if isinstance(value, tuple):
        safe_items, has_secret = _redact_mcp_config(list(value))
        return safe_items, has_secret
    return value, False


def _merge_mcp_config_secrets(current: Any, incoming: Any) -> Any:
    """Preserve persisted secrets when a public API placeholder is submitted."""

    if incoming == MCP_SECRET_PLACEHOLDER:
        return current
    if isinstance(incoming, dict):
        current_map = current if isinstance(current, dict) else {}
        return {
            str(key): _merge_mcp_config_secrets(current_map.get(str(key)), item)
            for key, item in incoming.items()
        }
    if isinstance(incoming, list):
        current_items = current if isinstance(current, list) else []
        return [
            _merge_mcp_config_secrets(
                current_items[index] if index < len(current_items) else None,
                item,
            )
            for index, item in enumerate(incoming)
        ]
    return incoming


def _artifact_tool_entrypoint() -> Path:
    configured = str(os.getenv("APP_ARTIFACT_TOOL_ENTRYPOINT") or "").strip()
    candidates = [
        Path(configured).expanduser() if configured else None,
        BASE_DIR / "node_modules" / "@oai" / "artifact-tool" / "dist" / "artifact_tool.mjs",
    ]
    for candidate in candidates:
        if candidate and candidate.is_file():
            return candidate.resolve()
    raise ToolError(
        "PowerPoint 生成组件尚未安装，请配置 APP_ARTIFACT_TOOL_ENTRYPOINT 后重试"
    )


def presentation_generation_status() -> dict[str, Any]:
    """Report whether the optional PPTX generator is ready on this host."""

    node_name = str(os.getenv("APP_NODE_BINARY") or "node").strip()
    node_binary = (
        str(Path(node_name).resolve())
        if Path(node_name).is_file()
        else shutil.which(node_name)
    )
    if not node_binary:
        return {"configured": False, "reason": "未找到 Node.js"}
    try:
        entrypoint = _artifact_tool_entrypoint()
    except ToolError as exc:
        return {"configured": False, "reason": str(exc)}
    if not PRESENTATION_GENERATOR_SCRIPT.is_file():
        return {"configured": False, "reason": "PowerPoint 生成脚本缺失"}
    return {"configured": True}


def _generate_pptx_with_artifact_tool(title: str, content: str, output_path: Path) -> dict[str, Any]:
    node_name = str(os.getenv("APP_NODE_BINARY") or "node").strip()
    node_binary = shutil.which(node_name) if not Path(node_name).is_file() else str(Path(node_name).resolve())
    if not node_binary:
        raise ToolError("PowerPoint 生成需要 Node.js，请安装后配置 APP_NODE_BINARY")
    if not PRESENTATION_GENERATOR_SCRIPT.is_file():
        raise ToolError("PowerPoint 生成脚本缺失，请联系平台管理员")

    try:
        timeout = max(10, min(int(os.getenv("APP_PRESENTATION_TIMEOUT_SECONDS", "180")), 600))
    except ValueError:
        timeout = 180
    payload = json.dumps({"title": title, "content": content}, ensure_ascii=False)
    command = [
        node_binary,
        str(PRESENTATION_GENERATOR_SCRIPT),
        "--artifact-tool",
        str(_artifact_tool_entrypoint()),
        "--output",
        str(output_path),
    ]
    try:
        completed = subprocess.run(
            command,
            input=payload,
            text=True,
            encoding="utf-8",
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        raise ToolError("PowerPoint 生成超时，请减少内容后重试") from exc
    except OSError as exc:
        raise ToolError("PowerPoint 生成服务无法启动，请联系平台管理员") from exc
    if completed.returncode != 0:
        technical_detail = (completed.stderr or completed.stdout or "unknown artifact-tool error")[-4000:]
        raise ToolError("PowerPoint 生成失败，请稍后重试或检查 Artifact Tool 运行环境") from RuntimeError(
            technical_detail
        )
    if not output_path.is_file() or output_path.stat().st_size < 4:
        raise ToolError("PowerPoint 生成失败：未得到可下载文件")
    with output_path.open("rb") as generated_file:
        if generated_file.read(4) != b"PK\x03\x04":
            raise ToolError("PowerPoint 生成失败：输出文件格式无效")
    try:
        return json.loads((completed.stdout or "{}").strip().splitlines()[-1])
    except (IndexError, json.JSONDecodeError):
        return {}


def resolve_artifact_path(relative_path: str | Path) -> Path:
    """Resolve an internal artifact path without allowing it to escape storage.

    Only relative, existing regular files below ``ARTIFACT_DIR`` are accepted.
    ``Path.resolve`` follows every existing symlink, so a link that points
    outside the configured artifact root is rejected by the containment check.
    """

    raw_value = str(relative_path or "").strip().replace("\\", "/")
    if (
        not raw_value
        or raw_value.startswith("/")
        or raw_value.startswith("//")
        or re.match(r"^[A-Za-z]:/", raw_value)
    ):
        raise ToolError("产物路径必须是受控目录内的相对路径")
    candidate_relative = Path(raw_value)
    if candidate_relative.is_absolute() or ".." in candidate_relative.parts:
        raise ToolError("产物路径不能包含目录穿越")

    try:
        root = ARTIFACT_DIR.resolve(strict=True)
        candidate = (root / candidate_relative).resolve(strict=True)
        candidate.relative_to(root)
    except FileNotFoundError as exc:
        raise ToolError("产物文件不存在") from exc
    except (OSError, RuntimeError, ValueError) as exc:
        raise ToolError("产物路径超出受控目录") from exc
    if not candidate.is_file():
        raise ToolError("产物路径不是文件")
    return candidate


def _tool(
    name: str,
    description: str,
    input_schema: dict[str, Any],
    *,
    effect: str,
) -> dict[str, Any]:
    """Describe a built-in tool and its side-effect class.

    Runtime permissions live outside ``McpGateway``. The gateway exposes
    immutable capability metadata so concurrent executions can evaluate their
    own permission snapshots. Built-ins use ``effect``; remote MCP tools use
    protocol ``annotations``.
    """

    return {
        "name": name,
        "description": description,
        "input_schema": input_schema,
        "effect": effect,
        "annotations": {"readOnlyHint": effect == "read"},
    }


CORE_BUILTIN_SERVERS = [
    {
        "id": "weather",
        "name": "天气预报 MCP",
        "kind": "builtin",
        "description": "通过 Open-Meteo 获取城市定位和结构化天气预报，无需 API Key。",
        "tools": [
            _tool("forecast", "查询指定城市今天、明天或后天的结构化天气预报。", {"type": "object", "properties": {"city": {"type": "string"}, "day": {"type": "string", "enum": ["today", "tomorrow", "day_after_tomorrow"], "default": "tomorrow"}}, "required": ["city"]}, effect="read"),
        ],
    },
    {
        "id": "web-search",
        "name": "联网搜索 MCP",
        "kind": "builtin",
        "description": "通过 Tavily 或 Brave Search 联网检索。默认需配置 API Key，未配置时能力状态会显示不可用。",
        "tools": [
            _tool("search", "搜索互联网并返回标题、链接与摘要。", {"type": "object", "properties": {"query": {"type": "string"}, "max_results": {"type": "integer", "default": 5}}, "required": ["query"]}, effect="read"),
        ],
    },
    {
        "id": "spreadsheet",
        "name": "表格生成 MCP",
        "kind": "builtin",
        "description": "生成 Excel 或 CSV 产物。",
        "tools": [
            _tool("create_excel", "根据行数据生成 XLSX 或 CSV 文件，输出格式由 filename 扩展名决定。", {"type": "object", "properties": {"rows": {"type": "array", "items": {"type": "object"}}, "filename": {"type": "string", "description": "以 .xlsx 或 .csv 结尾的文件名"}}, "required": ["rows", "filename"]}, effect="write"),
        ],
    },
    {
        "id": "report",
        "name": "报告生成 MCP",
        "kind": "builtin",
        "description": "生成 Word、PDF、PowerPoint、Markdown 和 HTML 文档。",
        "tools": [
            _tool("generate_markdown_report", "根据摘要和结果行生成 Markdown 报告文件。", {"type": "object", "properties": {"summary": {"type": "string"}, "rows": {"type": "array"}, "filename": {"type": "string"}}, "required": ["summary", "rows", "filename"]}, effect="write"),
            _tool("generate_document", "生成 DOCX、PDF、PPTX、Markdown 或 HTML 文档。", {"type": "object", "properties": {"title": {"type": "string"}, "content": {"type": "string"}, "format": {"type": "string", "enum": ["docx", "pdf", "pptx", "md", "html"]}, "filename": {"type": "string"}}, "required": ["title", "content", "format"]}, effect="write"),
        ],
    },
]

BUILTIN_SERVERS = CORE_BUILTIN_SERVERS


def server_to_runtime(row: dict[str, Any]) -> dict[str, Any]:
    return {
        **row,
        "enabled": bool(row.get("enabled")),
        "config": db.json_loads(row.get("config_json"), {}),
        "tools": db.json_loads(row.get("tools_json"), []),
    }


def server_to_api(server: dict[str, Any]) -> dict[str, Any]:
    """Serialize an MCP server without exposing credentials or raw JSON."""

    safe = {
        key: value
        for key, value in server.items()
        if key not in {"config_json", "tools_json"}
    }
    config, has_secret = _redact_mcp_config(server.get("config", {}))
    return {
        **safe,
        "enabled": bool(server.get("enabled")),
        "config": config,
        "tools": server.get("tools", []),
        "has_secret": has_secret,
        "secret_placeholder": MCP_SECRET_PLACEHOLDER,
    }


class McpGateway:
    def seed_builtin_servers(self) -> None:
        for server in BUILTIN_SERVERS:
            exists = db.query_one("SELECT id, kind FROM mcp_servers WHERE id = ?", (server["id"],))
            if exists:
                if exists.get("kind") == "builtin":
                    db.execute(
                        "UPDATE mcp_servers SET name = ?, description = ?, tools_json = ?, updated_at = ? WHERE id = ?",
                        (server["name"], server["description"], db.json_dumps(server["tools"]), db.utc_now(), server["id"]),
                    )
                continue
            now = db.utc_now()
            db.execute(
                """
                INSERT INTO mcp_servers(id, name, kind, description, enabled, config_json, tools_json, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    server["id"], server["name"], server["kind"], server["description"], 1,
                    db.json_dumps({}), db.json_dumps(server["tools"]), now, now,
                ),
            )

    def _list_servers_runtime(self, enabled_only: bool = False) -> list[dict[str, Any]]:
        sql = "SELECT * FROM mcp_servers"
        params: tuple[Any, ...] = ()
        if enabled_only:
            sql += " WHERE enabled = 1"
        sql += " ORDER BY name"
        return [server_to_runtime(r) for r in db.query_all(sql, params)]

    def _get_server_runtime(self, server_id: str) -> dict[str, Any] | None:
        row = db.query_one("SELECT * FROM mcp_servers WHERE id = ?", (server_id,))
        return server_to_runtime(row) if row else None

    @staticmethod
    def to_api(server: dict[str, Any]) -> dict[str, Any]:
        return server_to_api(server)

    def list_servers(self, enabled_only: bool = False) -> list[dict[str, Any]]:
        return [
            self.to_api(server)
            for server in self._list_servers_runtime(enabled_only=enabled_only)
        ]

    def get_server(self, server_id: str) -> dict[str, Any] | None:
        server = self._get_server_runtime(server_id)
        return self.to_api(server) if server else None

    def server_exists(self, server_id: str) -> bool:
        return self._get_server_runtime(server_id) is not None

    def _create_server_runtime(self, payload: dict[str, Any]) -> dict[str, Any]:
        now = db.utc_now()
        db.execute(
            """
            INSERT INTO mcp_servers(id, name, kind, description, enabled, config_json, tools_json, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                payload["id"], payload["name"], payload.get("kind", "http"), payload.get("description", ""),
                1 if payload.get("enabled", True) else 0, db.json_dumps(payload.get("config", {})),
                db.json_dumps(payload.get("tools", [])), now, now,
            ),
        )
        return self._get_server_runtime(payload["id"]) or payload

    def create_server(self, payload: dict[str, Any]) -> dict[str, Any]:
        return self.to_api(self._create_server_runtime(payload))

    def import_config(self, payload: Any) -> list[dict[str, Any]]:
        """Import platform MCP records or the common {mcpServers: {...}} format."""
        candidates: list[dict[str, Any]] = []
        if isinstance(payload, dict) and isinstance(payload.get("mcpServers"), dict):
            for server_id, raw_config in payload["mcpServers"].items():
                config = raw_config if isinstance(raw_config, dict) else {}
                kind = "mcp_http" if config.get("url") or config.get("base_url") else "mcp_stdio"
                candidates.append({
                    "id": str(server_id), "name": str(config.get("name") or server_id), "kind": kind,
                    "description": str(config.get("description") or "通过配置文件导入"), "enabled": bool(config.get("enabled", True)),
                    "config": {k: v for k, v in config.items() if k not in {"name", "description", "enabled", "tools"}},
                    "tools": config.get("tools") if isinstance(config.get("tools"), list) else [],
                })
        elif isinstance(payload, list):
            candidates = [x for x in payload if isinstance(x, dict)]
        elif isinstance(payload, dict):
            candidates = [payload]
        else:
            raise ValueError("MCP 配置必须是 JSON 对象或数组")

        imported: list[dict[str, Any]] = []
        for item in candidates:
            raw_id = str(item.get("id") or "").strip()
            raw = {
                "id": raw_id,
                "name": str(item.get("name") or raw_id),
                "kind": str(item.get("kind") or ("mcp_http" if item.get("url") else "mcp_stdio")),
                "description": str(item.get("description") or "通过配置文件导入"),
                "enabled": bool(item.get("enabled", True)),
                "config": item.get("config") if isinstance(item.get("config"), dict) else {
                    k: item[k] for k in ("command", "args", "env", "url", "headers") if k in item
                },
                "tools": item.get("tools") if isinstance(item.get("tools"), list) else [],
            }
            try:
                normalized = McpServerCreate.model_validate(raw).model_dump()
            except ValueError as exc:
                raise ValueError(f"MCP 配置无效：{exc}") from exc
            server_id = normalized["id"]
            if self._get_server_runtime(server_id):
                imported.append(self._update_server_runtime(server_id, normalized) or normalized)
            else:
                imported.append(self._create_server_runtime(normalized))
        if not imported:
            raise ValueError("配置中没有可导入的 MCP Server")
        return [self.to_api(server) for server in imported]

    def list_tools(self, server_id: str | None = None) -> list[dict[str, Any]]:
        servers = (
            [self._get_server_runtime(server_id)]
            if server_id
            else self._list_servers_runtime(enabled_only=True)
        )
        tools: list[dict[str, Any]] = []
        for server in servers:
            if not server:
                continue
            for tool in server.get("tools", []):
                tools.append({**tool, "server_id": server["id"], "server_name": server["name"]})
        return tools

    def get_tool_definition(self, server_id: str, tool_name: str) -> dict[str, Any] | None:
        """Return capability metadata without attaching execution state.

        A fresh dictionary is returned on every call.  This is important for
        parallel expert members: annotations may be inspected concurrently,
        but permissions must never be written onto the shared gateway or the
        persisted MCP server definition.
        """

        server = self._get_server_runtime(server_id)
        if not server:
            return None
        tool = next(
            (item for item in server.get("tools", []) if str(item.get("name") or "") == tool_name),
            None,
        )
        if not isinstance(tool, dict):
            return None
        return {
            **tool,
            "server_id": server_id,
            "server_name": server.get("name", server_id),
            "server_kind": server.get("kind", ""),
        }

    async def invoke_tool(self, server_id: str, tool_name: str, arguments: dict[str, Any], task_id: str | None = None) -> dict[str, Any]:
        server = self._get_server_runtime(server_id)
        if not server:
            raise ToolError(f"MCP server not found: {server_id}")
        if not server["enabled"]:
            raise ToolError(f"MCP server disabled: {server_id}")
        if server["kind"] == "builtin":
            return await self._invoke_builtin(server_id, tool_name, arguments, task_id=task_id)
        if server["kind"] == "http":
            return await self._invoke_http_tool(server, tool_name, arguments)
        if server["kind"] in {"mcp_stdio", "stdio"}:
            return await self._invoke_mcp_stdio(server, tool_name, arguments)
        if server["kind"] in {"mcp_http", "streamable_http"}:
            return await self._invoke_mcp_http(server, tool_name, arguments)
        raise ToolError(f"Unsupported MCP server kind: {server['kind']}")

    async def _invoke_builtin(self, server_id: str, tool_name: str, arguments: dict[str, Any], task_id: str | None = None) -> dict[str, Any]:
        if server_id == "web-search" and tool_name == "search":
            return await self._web_search(arguments)
        if server_id == "weather" and tool_name == "forecast":
            return await self._weather_forecast(arguments)
        if server_id == "spreadsheet" and tool_name == "create_excel":
            return await asyncio.to_thread(self._create_excel, arguments, task_id=task_id)
        if server_id == "report" and tool_name == "generate_markdown_report":
            return await asyncio.to_thread(self._generate_markdown_report, arguments, task_id=task_id)
        if server_id == "report" and tool_name == "generate_document":
            return await asyncio.to_thread(self._generate_document, arguments, task_id=task_id)
        raise ToolError(f"Builtin tool not found: {server_id}.{tool_name}")

    async def _invoke_http_tool(self, server: dict[str, Any], tool_name: str, arguments: dict[str, Any]) -> dict[str, Any]:
        tool = next((t for t in server.get("tools", []) if t.get("name") == tool_name), None)
        if not tool:
            raise ToolError(f"HTTP tool not found: {tool_name}")
        require_outbound_network("HTTP 工具调用", error_type=ToolError)
        if not env_flag("APP_ALLOW_HTTP_TOOLS"):
            raise ToolError("HTTP 工具尚未开启，请由管理员设置 APP_ALLOW_HTTP_TOOLS=true 后重启平台")
        config = server.get("config", {})
        base_url = config.get("base_url", "").rstrip("/")
        path = tool.get("path", "")
        method = tool.get("method", "POST").upper()
        url = base_url + path
        url = self._validate_remote_url(url)
        headers = self._resolve_env_values(config.get("headers", {}))
        timeout = float(config.get("timeout", 30))
        async with httpx.AsyncClient(timeout=timeout, follow_redirects=False) as client:
            if method == "GET":
                response = await client.get(url, params=arguments, headers=headers)
            else:
                response = await client.request(method, url, json=arguments, headers=headers)
            response.raise_for_status()
            try:
                return response.json()
            except Exception:
                return {"text": response.text}

    async def discover_tools(self, server_id: str) -> list[dict[str, Any]]:
        server = self._get_server_runtime(server_id)
        if not server:
            raise ToolError(f"MCP server not found: {server_id}")
        kind = server["kind"]
        if kind in {"builtin", "http"}:
            return self.list_tools(server_id)
        if kind in {"mcp_stdio", "stdio"}:
            tools = await self._list_mcp_stdio(server)
        elif kind in {"mcp_http", "streamable_http"}:
            tools = await self._list_mcp_http(server)
        else:
            raise ToolError(f"Unsupported MCP server kind: {kind}")
        db.execute("UPDATE mcp_servers SET tools_json = ?, updated_at = ? WHERE id = ?", (db.json_dumps(tools), db.utc_now(), server_id))
        return [{**tool, "server_id": server_id, "server_name": server["name"]} for tool in tools]

    def _update_server_runtime(self, server_id: str, payload: dict[str, Any]) -> dict[str, Any] | None:
        current = self._get_server_runtime(server_id)
        if not current:
            return None
        changes = {k: v for k, v in payload.items() if v is not None}
        if "config" in changes:
            changes["config"] = _merge_mcp_config_secrets(
                current.get("config", {}), changes["config"]
            )
        updated = {**current, **changes}
        db.execute(
            "UPDATE mcp_servers SET name = ?, kind = ?, description = ?, enabled = ?, config_json = ?, tools_json = ?, updated_at = ? WHERE id = ?",
            (updated["name"], updated["kind"], updated["description"], 1 if updated.get("enabled") else 0,
             db.json_dumps(updated.get("config", {})), db.json_dumps(updated.get("tools", [])), db.utc_now(), server_id),
        )
        return self._get_server_runtime(server_id)

    def update_server(self, server_id: str, payload: dict[str, Any]) -> dict[str, Any] | None:
        updated = self._update_server_runtime(server_id, payload)
        return self.to_api(updated) if updated else None

    def _resolve_env_values(self, value: Any) -> Any:
        if isinstance(value, str) and value.startswith("${") and value.endswith("}"):
            return os.getenv(value[2:-1], "")
        if isinstance(value, dict):
            return {k: self._resolve_env_values(v) for k, v in value.items()}
        if isinstance(value, list):
            return [self._resolve_env_values(v) for v in value]
        return value

    def _validate_remote_url(self, url: str) -> str:
        return validate_outbound_http_url(
            url,
            capability="MCP/HTTP",
            allowlist_env="APP_REMOTE_HOST_ALLOWLIST",
            require_allowlist=True,
            error_type=ToolError,
        )

    def _stdio_params(self, server: dict[str, Any]):
        if not env_flag("APP_ALLOW_STDIO_MCP"):
            raise ToolError("本地 stdio MCP 默认关闭，请设置 APP_ALLOW_STDIO_MCP=true")
        config = server.get("config", {})
        command = self._resolve_stdio_command(str(config.get("command", "")).strip())
        try:
            from mcp import StdioServerParameters
        except ImportError as exc:
            raise ToolError("未安装官方 MCP SDK，请执行 pip install -r requirements.txt") from exc
        return StdioServerParameters(command=command, args=[str(x) for x in config.get("args", [])], env=self._resolve_env_values(config.get("env", {})) or None)

    @staticmethod
    def _resolve_stdio_command(command: str) -> str:
        allowed = [
            item.strip()
            for item in os.getenv("APP_STDIO_COMMAND_ALLOWLIST", "").split(",")
            if item.strip()
        ]
        if not allowed:
            raise ToolError("本地 stdio MCP 要求管理员配置非空的 APP_STDIO_COMMAND_ALLOWLIST")
        if not command:
            raise ToolError("stdio command 不能为空")

        command_path = Path(command)
        if command_path.is_absolute():
            absolute_entries = [Path(item) for item in allowed if Path(item).is_absolute()]
            try:
                resolved = command_path.resolve(strict=True)
            except (OSError, RuntimeError) as exc:
                raise ToolError("stdio command 指向的可执行文件不存在") from exc
            allowed_match = False
            for entry in absolute_entries:
                try:
                    if resolved == entry.resolve(strict=True):
                        allowed_match = True
                        break
                except (OSError, RuntimeError):
                    continue
            if not allowed_match:
                raise ToolError("stdio 绝对路径不在 APP_STDIO_COMMAND_ALLOWLIST 中")
            if not resolved.is_file() or not os.access(resolved, os.X_OK):
                raise ToolError("stdio command 不是可执行文件")
            return str(resolved)

        if command_path.name != command or "/" in command or "\\" in command:
            raise ToolError("stdio command 只能使用裸命令或白名单中的绝对路径")
        bare_entries = [item for item in allowed if Path(item).name == item and "/" not in item and "\\" not in item]
        if command not in bare_entries:
            raise ToolError("stdio 裸命令不在 APP_STDIO_COMMAND_ALLOWLIST 中")
        resolved_name = shutil.which(command)
        if not resolved_name:
            raise ToolError("stdio command 在当前 PATH 中不可用")
        resolved = Path(resolved_name).resolve()
        if not resolved.is_file() or not os.access(resolved, os.X_OK):
            raise ToolError("stdio command 不是可执行文件")
        return str(resolved)

    @staticmethod
    def _normalize_annotations(value: Any) -> dict[str, Any]:
        if value is None:
            return {}
        if isinstance(value, dict):
            return dict(value)
        if hasattr(value, "model_dump"):
            try:
                dumped = value.model_dump(by_alias=True, exclude_none=True)
            except TypeError:
                dumped = value.model_dump()
            return dict(dumped) if isinstance(dumped, dict) else {}
        result: dict[str, Any] = {}
        for key in ("title", "readOnlyHint", "destructiveHint", "idempotentHint", "openWorldHint"):
            if hasattr(value, key):
                item = getattr(value, key)
                if item is not None:
                    result[key] = item
        return result

    def _normalize_mcp_tools(self, result: Any) -> list[dict[str, Any]]:
        tools: list[dict[str, Any]] = []
        for item in result.tools:
            tools.append(
                {
                    "name": item.name,
                    "description": item.description or "",
                    "input_schema": item.inputSchema or {},
                    "annotations": self._normalize_annotations(getattr(item, "annotations", None)),
                }
            )
        return tools

    def _normalize_mcp_result(self, result: Any) -> dict[str, Any]:
        content = []
        for item in getattr(result, "content", []) or []:
            if hasattr(item, "model_dump"):
                content.append(item.model_dump())
            else:
                content.append({"type": getattr(item, "type", "text"), "text": getattr(item, "text", str(item))})
        if bool(getattr(result, "isError", False)):
            messages = [
                str(item.get("text") or "").strip()
                for item in content
                if isinstance(item, dict) and str(item.get("text") or "").strip()
            ]
            detail = "；".join(messages)[:2_000] or "远程工具返回失败状态"
            raise ToolError(detail)
        return {
            "content": content,
            "is_error": False,
            "structured_content": getattr(result, "structuredContent", None),
        }

    async def _list_mcp_stdio(self, server: dict[str, Any]) -> list[dict[str, Any]]:
        from mcp import ClientSession
        from mcp.client.stdio import stdio_client
        timeout = self._stdio_timeout_seconds(server)
        try:
            async with asyncio.timeout(timeout):
                async with stdio_client(self._stdio_params(server)) as (read, write):
                    async with ClientSession(read, write) as session:
                        await session.initialize()
                        return self._normalize_mcp_tools(await session.list_tools())
        except TimeoutError as exc:
            raise ToolError(f"MCP stdio 工具发现超过 {timeout:g} 秒，已终止子进程") from exc

    async def _invoke_mcp_stdio(self, server: dict[str, Any], tool_name: str, arguments: dict[str, Any]) -> dict[str, Any]:
        from mcp import ClientSession
        from mcp.client.stdio import stdio_client
        timeout = self._stdio_timeout_seconds(server)
        try:
            async with asyncio.timeout(timeout):
                async with stdio_client(self._stdio_params(server)) as (read, write):
                    async with ClientSession(read, write) as session:
                        await session.initialize()
                        return self._normalize_mcp_result(await session.call_tool(tool_name, arguments))
        except TimeoutError as exc:
            raise ToolError(f"MCP stdio 调用 {tool_name} 超过 {timeout:g} 秒，已终止子进程") from exc

    def _stdio_timeout_seconds(self, server: dict[str, Any]) -> float:
        config = server.get("config", {})
        try:
            value = float(config.get("timeout", config.get("startup_timeout", 60)))
        except (TypeError, ValueError):
            value = 60.0
        return max(5.0, min(value, 300.0))

    def _mcp_http_config(self, server: dict[str, Any]) -> tuple[str, dict[str, str]]:
        require_outbound_network("远程 MCP", error_type=ToolError)
        if not env_flag("APP_ALLOW_REMOTE_MCP"):
            raise ToolError("远程 MCP 尚未开启，请由管理员设置 APP_ALLOW_REMOTE_MCP=true 后重启平台")
        config = server.get("config", {})
        url = str(config.get("url") or config.get("base_url") or "")
        url = self._validate_remote_url(url)
        return url, self._resolve_env_values(config.get("headers", {}))

    async def _list_mcp_http(self, server: dict[str, Any]) -> list[dict[str, Any]]:
        from mcp import ClientSession
        from mcp.client.streamable_http import streamable_http_client
        url, headers = self._mcp_http_config(server)
        async with httpx.AsyncClient(headers=headers, timeout=60, follow_redirects=False) as client:
            async with streamable_http_client(url, http_client=client) as (read, write, _):
                async with ClientSession(read, write) as session:
                    await session.initialize()
                    return self._normalize_mcp_tools(await session.list_tools())

    async def _invoke_mcp_http(self, server: dict[str, Any], tool_name: str, arguments: dict[str, Any]) -> dict[str, Any]:
        from mcp import ClientSession
        from mcp.client.streamable_http import streamable_http_client
        url, headers = self._mcp_http_config(server)
        async with httpx.AsyncClient(headers=headers, timeout=60, follow_redirects=False) as client:
            async with streamable_http_client(url, http_client=client) as (read, write, _):
                async with ClientSession(read, write) as session:
                    await session.initialize()
                    return self._normalize_mcp_result(await session.call_tool(tool_name, arguments))

    async def _web_search(self, arguments: dict[str, Any]) -> dict[str, Any]:
        query = str(arguments.get("query", "")).strip()
        if not query:
            raise ToolError("query 不能为空")
        require_outbound_network("联网搜索", error_type=ToolError)
        if not env_flag("APP_ALLOW_WEB_SEARCH"):
            raise ToolError("联网搜索尚未开启，请由管理员设置 APP_ALLOW_WEB_SEARCH=true 并配置搜索 API Key")
        limit = max(1, min(int(arguments.get("max_results", 5)), 10))
        tavily_key = os.getenv("TAVILY_API_KEY", "")
        brave_key = os.getenv("BRAVE_SEARCH_API_KEY", "")
        async with httpx.AsyncClient(timeout=30) as client:
            if tavily_key:
                response = await client.post("https://api.tavily.com/search", json={"api_key": tavily_key, "query": query, "max_results": limit})
                response.raise_for_status()
                data = response.json()
                return {"query": query, "provider": "tavily", "results": [{"title": r.get("title"), "url": r.get("url"), "snippet": r.get("content")} for r in data.get("results", [])]}
            if brave_key:
                response = await client.get("https://api.search.brave.com/res/v1/web/search", params={"q": query, "count": limit}, headers={"X-Subscription-Token": brave_key, "Accept": "application/json"})
                response.raise_for_status()
                data = response.json()
                return {"query": query, "provider": "brave", "results": [{"title": r.get("title"), "url": r.get("url"), "snippet": r.get("description")} for r in data.get("web", {}).get("results", [])]}
        raise ToolError("联网搜索尚未配置，请设置 TAVILY_API_KEY 或 BRAVE_SEARCH_API_KEY")

    async def _weather_forecast(self, arguments: dict[str, Any]) -> dict[str, Any]:
        city = str(arguments.get("city") or "").strip()
        if not city:
            raise ToolError("city 不能为空")
        require_outbound_network("天气查询", error_type=ToolError)
        day = str(arguments.get("day") or "tomorrow")
        day_index = {"today": 0, "tomorrow": 1, "day_after_tomorrow": 2}.get(day, 1)
        async with httpx.AsyncClient(timeout=30, headers={"User-Agent": "AgentNexus/0.1"}) as client:
            geo_response = await client.get(
                "https://geocoding-api.open-meteo.com/v1/search",
                params={"name": city, "count": 5, "language": "zh", "format": "json"},
            )
            geo_response.raise_for_status()
            candidates = geo_response.json().get("results") or []
            if not candidates:
                raise ToolError(f"未找到城市或地区：{city}")
            location = next((item for item in candidates if item.get("country_code") == "CN"), candidates[0])
            timezone = location.get("timezone") or "auto"
            forecast_response = await client.get(
                "https://api.open-meteo.com/v1/forecast",
                params={
                    "latitude": location["latitude"],
                    "longitude": location["longitude"],
                    "daily": "weather_code,temperature_2m_max,temperature_2m_min,precipitation_probability_max,precipitation_sum,wind_speed_10m_max,wind_gusts_10m_max",
                    "timezone": timezone,
                    "forecast_days": 3,
                },
            )
            forecast_response.raise_for_status()
            payload = forecast_response.json()
        daily = payload.get("daily") or {}
        dates = daily.get("time") or []
        if len(dates) <= day_index:
            raise ToolError("天气服务未返回所需日期的预报")

        def value(key: str) -> Any:
            values = daily.get(key) or []
            return values[day_index] if len(values) > day_index else None

        code = int(value("weather_code") or 0)
        weather_names = {
            0: "晴", 1: "大部晴朗", 2: "多云", 3: "阴",
            45: "有雾", 48: "雾凇", 51: "小毛毛雨", 53: "毛毛雨", 55: "强毛毛雨",
            61: "小雨", 63: "中雨", 65: "大雨", 66: "冻雨", 67: "强冻雨",
            71: "小雪", 73: "中雪", 75: "大雪", 77: "米雪",
            80: "小阵雨", 81: "阵雨", 82: "强阵雨", 85: "小阵雪", 86: "强阵雪",
            95: "雷暴", 96: "雷暴伴小冰雹", 99: "强雷暴伴冰雹",
        }
        return {
            "city": location.get("name") or city,
            "region": " ".join(x for x in [location.get("country"), location.get("admin1"), location.get("admin2")] if x),
            "date": dates[day_index],
            "day": day,
            "condition": weather_names.get(code, f"天气代码 {code}"),
            "weather_code": code,
            "temperature_max_c": value("temperature_2m_max"),
            "temperature_min_c": value("temperature_2m_min"),
            "precipitation_probability_max_percent": value("precipitation_probability_max"),
            "precipitation_sum_mm": value("precipitation_sum"),
            "wind_speed_max_kmh": value("wind_speed_10m_max"),
            "wind_gusts_max_kmh": value("wind_gusts_10m_max"),
            "timezone": payload.get("timezone"),
            "provider": "Open-Meteo",
            "source": "https://open-meteo.com/",
        }

    def _artifact_context(self, task_id: str | None) -> dict[str, str]:
        owner_task_id = str(task_id or "manual")
        workspace_id = "default"
        run_id = ""
        if task_id:
            task = db.query_one("SELECT workspace FROM tasks WHERE id = ?", (task_id,))
            if task:
                workspace_id = str(task.get("workspace") or "default")
            has_runs = db.query_one(
                "SELECT name FROM sqlite_master WHERE type = 'table' AND name = 'task_runs'"
            )
            if has_runs:
                run = db.query_one(
                    """
                    SELECT id FROM task_runs
                    WHERE task_id = ?
                    ORDER BY CASE
                        WHEN status IN ('running', 'paused', 'waiting_approval') THEN 0
                        WHEN status = 'queued' THEN 1
                        ELSE 2
                    END, attempt DESC
                    LIMIT 1
                    """,
                    (task_id,),
                )
                if run:
                    run_id = str(run.get("id") or "")
        return {
            "task_id": owner_task_id,
            "run_id": run_id,
            "workspace_id": workspace_id,
        }

    @staticmethod
    def _storage_segment(value: str, fallback: str) -> str:
        candidate = str(value or "").strip()
        if (
            candidate not in {"", ".", ".."}
            and len(candidate) <= 128
            and re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._-]*", candidate)
        ):
            return candidate
        digest = hashlib.sha256(candidate.encode("utf-8")).hexdigest()[:12]
        return f"{fallback}-{digest}"

    def _new_artifact_target(
        self, task_id: str | None, filename: str
    ) -> tuple[dict[str, str], str, Path]:
        context = self._artifact_context(task_id)
        artifact_id = "art_" + uuid.uuid4().hex[:12]
        task_segment = self._storage_segment(context["task_id"], "task")
        run_segment = self._storage_segment(context["run_id"], "unscoped")
        root = ARTIFACT_DIR.resolve(strict=False)
        output_dir = root / task_segment / run_segment / artifact_id
        output_dir.mkdir(parents=True, exist_ok=False)
        return context, artifact_id, output_dir / filename

    def _safe_filename(self, value: str, fallback: str) -> str:
        name = Path(value or fallback).name
        cleaned = "".join(c if c.isalnum() or c in "._-" else "_" for c in name)
        cleaned = cleaned[:160]
        return fallback if cleaned in {"", ".", ".."} else cleaned

    def _register_artifact(
        self,
        context: dict[str, str],
        artifact_id: str,
        name: str,
        kind: str,
        path: Path,
    ) -> dict[str, Any]:
        try:
            root = ARTIFACT_DIR.resolve(strict=True)
            resolved = path.resolve(strict=True)
            relative_path = resolved.relative_to(root).as_posix()
        except FileNotFoundError as exc:
            raise ToolError("产物文件不存在，无法登记") from exc
        except (OSError, RuntimeError, ValueError) as exc:
            raise ToolError("产物文件超出受控目录，拒绝登记") from exc
        resolved = resolve_artifact_path(relative_path)

        digest = hashlib.sha256()
        with resolved.open("rb") as artifact_file:
            for chunk in iter(lambda: artifact_file.read(1024 * 1024), b""):
                digest.update(chunk)
        size = resolved.stat().st_size
        mime_type = self._artifact_mime_type(name, kind)
        latest = db.query_one(
            "SELECT COALESCE(MAX(version), 0) AS max_version FROM artifacts WHERE task_id = ? AND name = ?",
            (context["task_id"], name),
        ) or {}
        version = int(latest.get("max_version") or 0) + 1
        metadata = {
            "storage": "immutable",
            "generator": "mcp_gateway",
        }
        db.execute(
            """
            INSERT INTO artifacts(
                id, task_id, run_id, workspace_id, name, kind, path,
                relative_path, mime_type, size, sha256, version,
                metadata_json, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                artifact_id,
                context["task_id"],
                context["run_id"],
                context["workspace_id"],
                name,
                kind,
                str(resolved),
                relative_path,
                mime_type,
                size,
                digest.hexdigest(),
                version,
                db.json_dumps(metadata),
                db.utc_now(),
            ),
        )
        return {
            "id": artifact_id,
            "name": name,
            "kind": kind,
            "run_id": context["run_id"],
            "workspace_id": context["workspace_id"],
            "relative_path": relative_path,
            "mime_type": mime_type,
            "size": size,
            "sha256": digest.hexdigest(),
            "version": version,
            "metadata": metadata,
            "download_url": f"/api/artifacts/{artifact_id}/download",
        }

    @staticmethod
    def _artifact_mime_type(name: str, kind: str) -> str:
        suffix = Path(name).suffix.lower()
        known = {
            ".csv": "text/csv",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".html": "text/html",
            ".md": "text/markdown",
            ".pdf": "application/pdf",
            ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }
        if suffix in known:
            return known[suffix]
        if kind == "markdown":
            return "text/markdown"
        return mimetypes.guess_type(name)[0] or "application/octet-stream"

    def _create_excel(self, arguments: dict[str, Any], task_id: str | None = None) -> dict[str, Any]:
        rows = arguments.get("rows") or []
        filename = self._safe_filename(arguments.get("filename") or "result.xlsx", "result.xlsx")
        suffix = Path(filename).suffix.lower()
        if suffix not in {".xlsx", ".csv"}:
            filename += ".xlsx"
            suffix = ".xlsx"
        context, artifact_id, out_path = self._new_artifact_target(task_id, filename)
        if rows:
            headers = list(rows[0].keys())
        else:
            headers = ["message"]
            rows = [{"message": "no data"}]
        if suffix == ".csv":
            with out_path.open("w", newline="", encoding="utf-8-sig") as stream:
                writer = csv.DictWriter(stream, fieldnames=headers, extrasaction="ignore")
                writer.writeheader()
                writer.writerows(rows)
            artifact = self._register_artifact(context, artifact_id, filename, "csv", out_path)
            return {"artifact": artifact, "row_count": len(rows), "format": "csv"}
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils import get_column_letter

            wb = Workbook()
            ws = wb.active
            ws.title = "results"
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="E8EEF7")
                cell.alignment = Alignment(horizontal="center")
            for row in rows:
                ws.append([row.get(h, "") for h in headers])
            for idx, header in enumerate(headers, start=1):
                width = max(len(str(header)) + 4, 16)
                for cell in ws[get_column_letter(idx)]:
                    width = min(max(width, len(str(cell.value or "")) + 4), 42)
                ws.column_dimensions[get_column_letter(idx)].width = width
            wb.save(out_path)
        except Exception:
            csv_path = out_path.with_suffix(".csv")
            with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
                writer = csv.DictWriter(f, fieldnames=headers, extrasaction="ignore")
                writer.writeheader()
                writer.writerows(rows)
            artifact = self._register_artifact(context, artifact_id, csv_path.name, "csv", csv_path)
            return {"artifact": artifact, "row_count": len(rows), "format": "csv"}
        artifact = self._register_artifact(context, artifact_id, filename, "xlsx", out_path)
        return {"artifact": artifact, "row_count": len(rows), "format": "xlsx"}

    def _generate_markdown_report(self, arguments: dict[str, Any], task_id: str | None = None) -> dict[str, Any]:
        summary = arguments.get("summary", "")
        rows = arguments.get("rows") or []
        filename = self._safe_filename(arguments.get("filename") or "report.md", "report.md")
        if not filename.endswith(".md"):
            filename += ".md"
        context, artifact_id, out_path = self._new_artifact_target(task_id, filename)
        lines = ["# 任务报告", "", summary, "", "## 结果明细", ""]
        if rows:
            headers = list(rows[0].keys())
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows[:50]:
                lines.append("| " + " | ".join(str(row.get(h, "")) for h in headers) + " |")
        else:
            lines.append("无明细数据。")
        lines.append("")
        lines.append("---")
        out_path.write_text("\n".join(lines), encoding="utf-8")
        artifact = self._register_artifact(context, artifact_id, filename, "markdown", out_path)
        return {"artifact": artifact, "row_count": len(rows)}

    def _generate_document(self, arguments: dict[str, Any], task_id: str | None = None) -> dict[str, Any]:
        title = str(arguments.get("title") or "任务文档")
        content = str(arguments.get("content") or "")
        fmt = str(arguments.get("format") or "md").lower()
        if fmt not in {"docx", "pdf", "pptx", "md", "html"}:
            raise ToolError("文档格式仅支持 docx、pdf、pptx、md、html")
        filename = self._safe_filename(arguments.get("filename") or f"document.{fmt}", f"document.{fmt}")
        if not filename.lower().endswith(f".{fmt}"):
            filename += f".{fmt}"
        context, artifact_id, out_path = self._new_artifact_target(task_id, filename)
        if fmt == "md":
            out_path.write_text(f"# {title}\n\n{content}\n", encoding="utf-8")
            kind = "markdown"
        elif fmt == "html":
            body: list[str] = []
            in_list = False
            for raw in content.splitlines():
                line = raw.strip()
                if not line:
                    if in_list:
                        body.append("</ul>")
                        in_list = False
                    continue
                heading = re.match(r"^(#{1,6})\s+(.+)$", line)
                if heading:
                    if in_list:
                        body.append("</ul>")
                        in_list = False
                    level = min(len(heading.group(1)) + 1, 6)
                    body.append(f"<h{level}>{html.escape(self._plain_markdown(heading.group(2)))}</h{level}>")
                elif re.match(r"^[-*]\s+", line):
                    if not in_list:
                        body.append("<ul>")
                        in_list = True
                    body.append(f"<li>{html.escape(self._plain_markdown(re.sub(r'^[-*]\\s+', '', line)))}</li>")
                else:
                    if in_list:
                        body.append("</ul>")
                        in_list = False
                    body.append(f"<p>{html.escape(self._plain_markdown(line))}</p>")
            if in_list:
                body.append("</ul>")
            safe_title = html.escape(title)
            out_path.write_text(
                "<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
                f"<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\"><title>{safe_title}</title>"
                "<style>body{max-width:920px;margin:40px auto;padding:0 24px;color:#202637;font:16px/1.75 system-ui,sans-serif}"
                "h1,h2,h3{color:#27245f}h1{border-bottom:2px solid #716bea;padding-bottom:12px}li{margin:6px 0}</style></head>"
                f"<body><h1>{safe_title}</h1>{''.join(body)}</body></html>",
                encoding="utf-8",
            )
            kind = "html"
        elif fmt == "pptx":
            _generate_pptx_with_artifact_tool(title, content, out_path)
            kind = "pptx"
        elif fmt == "docx":
            from docx import Document
            document = Document()
            document.add_heading(title, level=0)
            lines = content.splitlines()
            index = 0
            while index < len(lines):
                line = lines[index].strip()
                if not line:
                    index += 1
                    continue
                if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|\s*$", lines[index + 1]):
                    table_lines = [line]
                    index += 2
                    while index < len(lines) and lines[index].strip().startswith("|"):
                        table_lines.append(lines[index].strip())
                        index += 1
                    rows = [[self._plain_markdown(cell.strip()) for cell in row.strip("|").split("|")] for row in table_lines]
                    table = document.add_table(rows=1, cols=len(rows[0]))
                    table.style = "Table Grid"
                    for col, value in enumerate(rows[0]):
                        table.rows[0].cells[col].text = value
                    for values in rows[1:]:
                        cells = table.add_row().cells
                        for col, value in enumerate(values[:len(cells)]):
                            cells[col].text = value
                    continue
                heading = re.match(r"^(#{1,6})\s+(.+)$", line)
                if heading:
                    document.add_heading(self._plain_markdown(heading.group(2)), level=min(len(heading.group(1)), 4))
                elif re.match(r"^[-*]\s+", line):
                    document.add_paragraph(self._plain_markdown(re.sub(r"^[-*]\s+", "", line)), style="List Bullet")
                elif re.match(r"^\d+[.)]\s+", line):
                    document.add_paragraph(self._plain_markdown(re.sub(r"^\d+[.)]\s+", "", line)), style="List Number")
                else:
                    document.add_paragraph(self._plain_markdown(line))
                index += 1
            document.save(out_path)
            kind = "docx"
        else:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            styles = getSampleStyleSheet()
            styles["Title"].fontName = "STSong-Light"
            styles["BodyText"].fontName = "STSong-Light"
            story = [Paragraph(title.replace("&", "&amp;"), styles["Title"]), Spacer(1, 12)]
            for block in content.split("\n"):
                story.append(Paragraph(block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or " ", styles["BodyText"]))
                story.append(Spacer(1, 5))
            SimpleDocTemplate(str(out_path), pagesize=A4).build(story)
            kind = "pdf"
        return {
            "artifact": self._register_artifact(context, artifact_id, filename, kind, out_path),
            "format": fmt,
        }

    def _plain_markdown(self, text: str) -> str:
        return re.sub(r"[*_`]+", "", text).strip()
